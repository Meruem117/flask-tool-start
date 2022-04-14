from docx import Document
import os


def open_docx(dirname, filename):
    docx_path = os.path.join(dirname, filename)
    document = Document(docx_path)
    return document


def handle_table(docx):
    table = docx.tables[0]
    replace_text(table, 0, 2, 0, 0, '2022年3月13日15时59分')
    replace_text(table, 2, 2, 0, 0, '王守义')
    replace_text(table, 2, 2, 0, 1, '北港')
    docx.save('../static/tmp/demo.docx')


def replace_text(table, row, col, para, r, text):
    run = table.cell(row, col).paragraphs[para].runs[r]
    run.text = run.text.replace(run.text, text)


def add_text(table, row, col, para, text):
    run = table.cell(row, col).paragraphs[para]
    run.text = run.text.replace(run.text, text)


def map_doc_table(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        print(run.text)


if __name__ == '__main__':
    docx_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir)) + '\\static\\docx'
    doc = open_docx(docx_dir, 'doc1.docx')
    map_doc_table(doc)
    # handle_table(doc)
